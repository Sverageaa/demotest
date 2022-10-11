import json
import re
import time
import warnings

from io import BytesIO

import requests
from selenium import webdriver
from selenium.webdriver.common.by import By
import urllib.request

from lxml import etree
from docx import Document  # 用来建立一个word对象
from docx.shared import Pt, Cm  # 用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  # 设置字体
from docx.shared import RGBColor  # 设置字体的颜色
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对其方式

from selenium.webdriver.chrome.options import Options

import sqlite3

# global mekey
# global biaoti

# 通用文档变量
# global doc
global dates
sh = {}
dates = {}

global gfm
gfm = {}

warnings.filterwarnings("ignore", category=Warning)

chrome_options = webdriver.ChromeOptions()
chrome_options.add_argument('--headless')  # 增加无界面选项
chrome_options.add_argument('--disable-gpu')  # 如果不加这个选项，有时定位会出现问题
chrome_options.binary_location = "C:\\Users\\Administrator\\AppData\\Local\\Google\\Chrome\\Application\\chrome.exe"
browser = webdriver.Chrome(chrome_options=chrome_options)


def xhjm(xxuid, cookies, headers, strtime, dates):
    browser.get("http://wm.gongzuo966.com/do.php?b=" + strtime)

    dates['text'] = browser.find_element(By.XPATH,'//p').text
    dates['uuid'] = xxuid

    response = requests.post('https://www.gongbiaoku.com/api/reader/copy', cookies=cookies, headers=headers, data=dates)

    # 解码后的内容
    # print(response.text)

    html2 = json.loads(response.text)
    return html2


# 解码
def jm(strm, xxuid):
    global dates

    cookies = {
        'rememberMe': 'YKYXKXwEaRou/r7CKSnlwGMFNNoq18/q0MzJo4lRKZjtjKVStCeWgOj3aYHb7exJ7TM7yzPZHIaBaUrNGDDdJnoOq0kxVkWjgbSVXOg0TFX7wGece61PmQaaB3zlq6RFISLrenLldZcjuMUkkzVHxxMa48BOrfYV9jTdfJD8CiRxqR0C1Pur1J7c6Ogt8WbBhFYhLHgVcVVDtxyyCmqBxvE70H8Fmfbq2k2p76Qwj7CqJuyIc+K3TqFjC5Pn50BZpwyS6KQvwH4aQisp5FTPU7G+7CoacsfJoznP8T3gmZBcWc14p8bqgXSnHwMQJNDe66T1YLTlo8juhWzWjUhsa6s7hiGQU85kxVnIko0i92rTZVF5OWRJE5ItW2OeqgQ6wnmH6AMIHGDfCmogiOA1nrMQsBhtBG0cMxHSnPBiwTSb1/j7ZvMqE/JaSVuzpqI+O6OdqLxm+N3gA7UagkidhcgDmjQjqCySCHbmSnAnaGDHssTXyLUjaDyWEbG5ilRyud/EvgjaLU1BxPrqE5hdx543jvsgslIHO98C2aMQhL5qb1YwCthy3qK4vbZ/R2QvBLzA00J/6QHABjJI02nC1Sir',
        'JSESSIONID': 'B08BADCC09CDE541FD14E8C7AB9063C2',
        'JSESSIONID': 'B08BADCC09CDE541FD14E8C7AB9063C2',
        'clientlanguage': 'und',
    }

    headers = {
        'Connection': 'keep-alive',
        'Accept': '*/*',
        'X-Requested-With': 'XMLHttpRequest',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36',
        'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'Origin': 'https://www.gongbiaoku.com',
        'Sec-Fetch-Site': 'same-origin',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': 'https://www.gongbiaoku.com/read/a2z18467ck3?secId=44j80245592y',
        'Accept-Language': 'zh-CN,zh;q=0.9',
        # Requests sorts cookies= alphabetically
        # 'Cookie': 'rememberMe=YKYXKXwEaRou/r7CKSnlwGMFNNoq18/q0MzJo4lRKZjtjKVStCeWgOj3aYHb7exJ7TM7yzPZHIaBaUrNGDDdJnoOq0kxVkWjgbSVXOg0TFX7wGece61PmQaaB3zlq6RFISLrenLldZcjuMUkkzVHxxMa48BOrfYV9jTdfJD8CiRxqR0C1Pur1J7c6Ogt8WbBhFYhLHgVcVVDtxyyCmqBxvE70H8Fmfbq2k2p76Qwj7CqJuyIc+K3TqFjC5Pn50BZpwyS6KQvwH4aQisp5FTPU7G+7CoacsfJoznP8T3gmZBcWc14p8bqgXSnHwMQJNDe66T1YLTlo8juhWzWjUhsa6s7hiGQU85kxVnIko0i92rTZVF5OWRJE5ItW2OeqgQ6wnmH6AMIHGDfCmogiOA1nrMQsBhtBG0cMxHSnPBiwTSb1/j7ZvMqE/JaSVuzpqI+O6OdqLxm+N3gA7UagkidhcgDmjQjqCySCHbmSnAnaGDHssTXyLUjaDyWEbG5ilRyud/EvgjaLU1BxPrqE5hdx543jvsgslIHO98C2aMQhL5qb1YwCthy3qK4vbZ/R2QvBLzA00J/6QHABjJI02nC1Sir; JSESSIONID=B08BADCC09CDE541FD14E8C7AB9063C2; JSESSIONID=B08BADCC09CDE541FD14E8C7AB9063C2; clientlanguage=und',
    }

    # 切割一下长度 150  因为对方做了解码长度限制
    m = 0
    fstr = ""
    zdstr = []
    strm = strm.replace("'", '"')
    for i in range(len(strm)):
        if m > 150:
            m = 0
            fstr = fstr + strm[i]
            zdstr.append(fstr)
            fstr = ""

        else:
            m = m + 1
            fstr = fstr + strm[i]

    if fstr != "":
        zdstr.append(fstr)

    fanhui = ""
    for i in range(len(zdstr)):
        strtime = str(int(round(time.time() * 1000)))
        datap = {
            "b": strtime,
            "n": zdstr[i]
        }

        # 这两个链接 是浏览器解码用的
        aaa = requests.post('http://wm.gongzuo966.com/doq.php', data=datap)

        browser.get("http://wm.gongzuo966.com/do.php?b=" + strtime)

        dates['text'] = browser.find_element(By.XPATH,'//p').text
        dates['uuid'] = xxuid

        response = requests.post('https://www.gongbiaoku.com/api/reader/copy', cookies=cookies, headers=headers,
                                 data=dates)

        # 解码后的内容
        print(response.text)
        html2 = json.loads(response.text)

        # 这个位置加了返回失败判断，这个不一定必须加，返回失败不一定都会出现，屏蔽也可能是别原因，这个随后根据实际情况调吧，可以删除
        if html2['status'] == "error":
            # print("获取数据失败")

            kdsoaoa = ""
            # 我测试用的这里
            # print(dates)
            # kdsoaoa = html2['message']
            #
            # while True:
            #     time.sleep(1)
            #     mmaaaff = xhjm(xxuid, cookies, headers, strtime, dates)
            #
            #     if mmaaaff['status']=="success":
            #         kdsoaoa = mmaaaff['message']
            #         break
            #     print("获取数据失败")
        else:
            kdsoaoa = html2['message']
        fanhui = fanhui + kdsoaoa
    return fanhui


def get_val(uuid):
    global dates

    cookies = {
        'rememberMe': 'YKYXKXwEaRou/r7CKSnlwGMFNNoq18/q0MzJo4lRKZjtjKVStCeWgOj3aYHb7exJ7TM7yzPZHIaBaUrNGDDdJnoOq0kxVkWjgbSVXOg0TFX7wGece61PmQaaB3zlq6RFISLrenLldZcjuMUkkzVHxxMa48BOrfYV9jTdfJD8CiRxqR0C1Pur1J7c6Ogt8WbBhFYhLHgVcVVDtxyyCmqBxvE70H8Fmfbq2k2p76Qwj7CqJuyIc+K3TqFjC5Pn50BZpwyS6KQvwH4aQisp5FTPU7G+7CoacsfJoznP8T3gmZBcWc14p8bqgXSnHwMQJNDe66T1YLTlo8juhWzWjUhsa6s7hiGQU85kxVnIko0i92rTZVF5OWRJE5ItW2OeqgQ6wnmH6AMIHGDfCmogiOA1nrMQsBhtBG0cMxHSnPBiwTSb1/j7ZvMqE/JaSVuzpqI+O6OdqLxm+N3gA7UagkidhcgDmjQjqCySCHbmSnAnaGDHssTXyLUjaDyWEbG5ilRyud/EvgjaLU1BxPrqE5hdx543jvsgslIHO98C2aMQhL5qb1YwCthy3qK4vbZ/R2QvBLzA00J/6QHABjJI02nC1Sir',
        'JSESSIONID': 'B08BADCC09CDE541FD14E8C7AB9063C2',
        'JSESSIONID': 'B08BADCC09CDE541FD14E8C7AB9063C2',
        'clientlanguage': 'und',
    }

    headers = {
        'Connection': 'keep-alive',
        'Accept': '*/*',
        'X-Requested-With': 'XMLHttpRequest',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36',
        'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'Origin': 'https://www.gongbiaoku.com',
        'Sec-Fetch-Site': 'same-origin',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': 'https://www.gongbiaoku.com/read/a2z18467ck3?secId=44j80245592y',
        'Accept-Language': 'zh-CN,zh;q=0.9',
        # Requests sorts cookies= alphabetically
        # 'Cookie': 'rememberMe=YKYXKXwEaRou/r7CKSnlwGMFNNoq18/q0MzJo4lRKZjtjKVStCeWgOj3aYHb7exJ7TM7yzPZHIaBaUrNGDDdJnoOq0kxVkWjgbSVXOg0TFX7wGece61PmQaaB3zlq6RFISLrenLldZcjuMUkkzVHxxMa48BOrfYV9jTdfJD8CiRxqR0C1Pur1J7c6Ogt8WbBhFYhLHgVcVVDtxyyCmqBxvE70H8Fmfbq2k2p76Qwj7CqJuyIc+K3TqFjC5Pn50BZpwyS6KQvwH4aQisp5FTPU7G+7CoacsfJoznP8T3gmZBcWc14p8bqgXSnHwMQJNDe66T1YLTlo8juhWzWjUhsa6s7hiGQU85kxVnIko0i92rTZVF5OWRJE5ItW2OeqgQ6wnmH6AMIHGDfCmogiOA1nrMQsBhtBG0cMxHSnPBiwTSb1/j7ZvMqE/JaSVuzpqI+O6OdqLxm+N3gA7UagkidhcgDmjQjqCySCHbmSnAnaGDHssTXyLUjaDyWEbG5ilRyud/EvgjaLU1BxPrqE5hdx543jvsgslIHO98C2aMQhL5qb1YwCthy3qK4vbZ/R2QvBLzA00J/6QHABjJI02nC1Sir; JSESSIONID=B08BADCC09CDE541FD14E8C7AB9063C2; JSESSIONID=B08BADCC09CDE541FD14E8C7AB9063C2; clientlanguage=und',
    }

    response = requests.post(f'https://www.gongbiaoku.com/read/content/' + uuid, cookies=cookies, headers=headers)

    mm = response.text

    html = etree.HTML(response.text)

    # print(re.findall("font-family:'myFont(.*?)_", response.text))
    key = re.findall("font-family:'myFont(.*?)_", response.text)[0]

    keyword = ''.join(html.xpath('.//text()')[1:])
    name = keyword
    newname = ''
    for char in name:
        if char not in newname:
            newname += char

    dict = {}
    for i in range(0, int(len(newname) / 200 + 0.999999999)):
        dates = {
            'key': key,
            'text': '鄡璴恽夓',
            'uuid': uuid,
        }

    #  mekey = key

    ff = ''
    flat = {}

    def replace_up(mm, str1, str2):
        mm2 = ''
        for j in range(0, len(mm)):
            try:
                if (flat[j] == 1):
                    mm2 = mm2 + mm[j]
            except:
                if (mm[j] == str1):
                    flat[j] = 1
                    mm2 = mm2 + str2
                else:
                    mm2 = mm2 + mm[j]
        mm = mm2
        return mm

    for j in dict:
        if (j in mm):
            mm = replace_up(mm, j, dict[j])

    return mm


def get_val2(uuid):
    test = get_val(uuid).replace('看新旧标准变化', '').replace('添加对比', '')

    html = etree.HTML(test)

    for i in html.xpath('.//div[@class="item"]'):
        pass
        # print(i.xpath('.//div/text()'))
    test = test.replace('class="item"', 'class="item" style="color: blue"')
    test = test.replace(' class="bold"', ' class="bold" style="color: red"')
    # 去除笔记 条文说明
    for i in re.findall('<a class="xy scbtn collect(.*?)</a>', test):
        # print('<a class="xy scbtn collect' + i + '</a>')
        test = test.replace('<a class="xy scbtn collect' + i + '</a>', '')
    for i in re.findall('<a class="xy sharebtn(.*?)</a>', test):
        # print('<a class="xy sharebtn' + i + '</a>')
        test = test.replace('<a class="xy sharebtn' + i + '</a>', '')
    for i in re.findall('<a class="xy notebtn note(.*?)</a>', test):
        # print('<a class="xy notebtn note' + i + '</a>')
        test = test.replace('<a class="xy notebtn note' + i + '</a>', '')
    for i in re.findall('<a class="showbtn(.*?)</a>', test):
        pass
        # print('<a class="showbtn' + i + '</a>')
        # test = test.replace('<a class="showbtn'+i+'</a>', '')
    test = test.replace('style="margin-right: 5px;vertical-align: middle;margin-bottom: 3px;', 'style=" width="700"')
    return test


def get_val3(uuid):
    html = get_val2(uuid)

    # </h3><div style="display: inline;"> </div>
    html2 = html
    for i in re.findall('</h3><div style="display: inline;">(.*?)</div>', re.findall('<body>(.*?)</body>', html)[0]):
        ttt = '</h3><div style="display: inline;">' + i + '</div>'
        html2 = html2.replace(ttt, '</h3></div><p></p><span class="noFd">' + i + '</span>')
        html2 = html2.replace('<h3 style="display: inline;font-weight: bolder !important;">',
                              '<div style="display: inline;"><h3  style="display: inline;">')

    html3 = html2
    for i in re.findall('<div class="noFd">(.*?)</div>', html2):
        ttt = '<div class="noFd">' + i + '</div>'
        html3 = html3.replace(ttt, ttt)

    return re.findall('<body>(.*?)</body>', html3)[0]


# div 没有class的情况
def zhl(html2, biaoti, xxuid, doc):
    htmla = etree.HTML(html2)
    title = jm(str(htmla.xpath('//h2')[0].text), xxuid).replace("　", "").replace(" ", "").replace("\r\n", "")

    if str(title) != "None":
        if str(title) != biaoti.replace("\r\n", "").replace("　", ""):
            ja = doc.add_heading("", 2)
            ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ja.add_run(title)
            r.font.name = u'MS Gothic'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
            r.font.size = Pt(13)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    htmlx = etree.HTML(html2)
    divas = htmlx.xpath('//h3/div')

    for xmh in range(len(divas)):
        zhuan = str(etree.tostring(divas[xmh], encoding="utf-8", pretty_print=True, method="html").decode("utf-8"))

        response = etree.HTML(zhuan)

        tupians = response.xpath('//img[contains(@class, "ar_image_")]/@src')
        chun = jm(str(response.xpath('string(.)')).replace("\r\n", "").replace("\n", "").replace("           ", ""), xxuid)
        chun.replace("          ", "")
        chun.replace("         ", "")
        chun.replace("        ", "")
        chun.replace("       ", "")
        chun.replace("      ", "")
        chun.replace("     ", "")
        chun.replace("    ", "")
        chun.replace("   ", "")
        chun.replace("  ", "")

        if (len(tupians) > 0) and (tupians != "[]"):
            ja = doc.add_paragraph()
            r = ja.add_run(jm(str(chun), xxuid).replace("    ", ""))
            r.font.name = u'宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            r.font.size = Pt(11)
            ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r.bold = True

            for tuim in range(len(tupians)):
                response = requests.get(tupians[tuim])

                binary_img = BytesIO(response.content)

                imja = doc.add_picture(binary_img)
                imja.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            ja = doc.add_paragraph()
            r = ja.add_run(chun)
            r.font.name = u'宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            r.font.size = Pt(11)
            ja.alignment = WD_ALIGN_PARAGRAPH.LEFT


# div 有class的情况
def xiedoc(html2, biaoti, xxuid, doc):
    htmlx = etree.HTML(html2)
    divas = htmlx.xpath('//div[contains(@class,"item-content")]')

    if len(divas) < 1:
        zhl(html2, biaoti, xxuid, doc)
    else:
        htmla = etree.HTML(html2)
        title = jm(str(htmla.xpath('//h2')[0].text), xxuid).replace("　", "").replace(" ", "").replace("\r\n", "")

        if str(title) != "None":
            if str(title) != biaoti.replace("\r\n", "").replace("　", ""):
                ja = doc.add_heading("", 2)
                ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = ja.add_run(title)
                r.font.name = u'MS Gothic'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
                r.font.size = Pt(13)
                r.bold = True
                r.font.color.rgb = RGBColor(0, 0, 0)

        for xmh in range(len(divas)):
            # 两个好像重复了，不过为了方便就没改  下面就没改
            nr = etree.tostring(divas[xmh], encoding="utf-8", pretty_print=True, method="html").decode("utf-8")
            nrmmm = etree.tostring(divas[xmh], encoding="utf-8", pretty_print=True, method="html").decode("utf-8")

            # print("解析出来的内容aa")
            # print(nr)

            ja = doc.add_paragraph()
            r = ja.add_run("")
            r.font.name = u'宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            r.font.size = Pt(11)
            ja.alignment = WD_ALIGN_PARAGRAPH.LEFT

            r = ja.add_run("\n\n")
            r.font.name = u'宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            r.font.size = Pt(11)
            ja.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 直接区分是否存在span标签
            if "</span>" in nr:
                diyi = str(nr).split("</span>")

                response = etree.HTML(diyi[0])
                chun = jm(str(response.xpath('string(.)')).replace("\r\n", "").replace("\n", "").replace("      ", ""), xxuid)

                ja = doc.add_paragraph()
                r = ja.add_run(chun)
                r.font.name = u'宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                r.font.size = Pt(11)

                if 'class="item-content forceinfo show"' in nr.replace("\n", ""):
                    r.bold = True

                ja.alignment = WD_ALIGN_PARAGRAPH.LEFT
                diyin = ""

                for dii in range(len(diyi)):
                    if dii > 0:
                        diyin = diyin + diyi[dii]

                # print(diyin)
                shuoming = ""
                # 判断是否存在说明  存在着分开
                if 'class="clause"' in nr:

                    # 怎进行分割
                    fenge = diyin.split('class="clause"')
                    zhengwen = fenge[0]
                    shuoming = fenge[1]
                else:
                    zhengwen = nr

                rdim = etree.HTML(zhengwen)
                diaaaa = rdim.xpath('//div[@class="align-center"]')

                if len(diaaaa) < 1:
                    diaaaa = rdim.xpath('//div[@class="noFd"]')

                if (len(diaaaa) < 1) and (shuoming == ""):
                    chun = str(rdim.xpath('string(.)')).replace("\r\n", "").replace("\n", "")

                    jiemahou = jm(str(chun), xxuid).replace("    ", "")
                    # 直接追加 到后面
                    r = ja.add_run(jiemahou)
                    r.font.name = u'宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    r.font.size = Pt(11)

                    # 情况颇多，说不定某一个特殊的又出现另外的情况
                    tupians = rdim.xpath('//img[@class="ar_image_2"]/@src')

                    for tuim in range(len(tupians)):
                        response = requests.get(tupians[tuim])
                        binary_img = BytesIO(response.content)
                        imja = doc.add_picture(binary_img)
                        imja.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for diai in range(len(diaaaa)):
                    # 提取内容
                    nr = etree.tostring(diaaaa[diai], encoding="utf-8", pretty_print=True, method="html").decode(
                        "utf-8")

                    response = etree.HTML(nr)
                    tupians = response.xpath('//img[contains(@class, "ar_image_")]/@src')
                    chun = str(response.xpath('string(.)')).replace("\r\n", "").replace("\n", "")

                    if chun != "":
                        # 下面的一点是 你让加的代码
                        # chun = chun.replace(" ", "")

                        if "mm2" in chun:
                            chun = chun.replace("mm2", "mm²")
                        elif "m2" in chun:
                            chun = chun.replace("m2", "㎡")
                        elif "m3" in chun:
                            chun = chun.replace("m3", "m³")
                        elif "mm3" in chun:
                            tx = chun.replace("mm3", "mm³")
                        elif "m4" in chun:
                            tx = chun.replace("m4", "m⁴")
                        elif "hm2" in chun:
                            chun = chun.replace("hm2", "hm²")
                        elif "km2" in chun:
                            chun = chun.replace("km2", "km²")
                        else:
                            pass

                        # 这里区分是否存在图片
                        if (len(tupians) > 0) and (tupians != "[]"):
                            ja = doc.add_paragraph()
                            r = ja.add_run(jm(str(chun), xxuid).replace(" ", ""))
                            r.font.name = u'宋体'
                            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                            r.font.size = Pt(11)
                            ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r.bold = True

                            for tuim in range(len(tupians)):
                                response = requests.get(tupians[tuim])
                                binary_img = BytesIO(response.content)
                                imja = doc.add_picture(binary_img)
                                imja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            ja = doc.add_paragraph()
                            r = ja.add_run("\n" + jm(chun, xxuid))
                            r.font.name = u'宋体'
                            r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                            r.font.size = Pt(11)
                            # 字体加粗
                            if 'class="item-content forceinfo show"' in nrmmm.replace("\n", ""):
                                r.bold = True
                            ja.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # 说明内容
                if shuoming != "":
                    shuoming = shuoming.split('class="item"')[1]
                    ja = doc.add_paragraph()
                    r = ja.add_run("\n" + "显示条文说明")
                    r.font.name = u'宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    r.font.size = Pt(11)
                    ja.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # 提取说明内容
                    # diazs = rdim.xpath('//div[@class="clause"]')
                    # nr = etree.tostring(diazs[0], encoding="utf-8", pretty_print=True, method="html").decode("utf-8")
                    # divzs = str(shuoming).split('class="item"')[1]
                    # divze = re.findall(r'<div.*?</div>', shuoming)
                    shuoming = shuoming.replace('<div class="image-title">', '<p ').replace('<div class="image-desc', '<p ')
                    reshuoming = etree.HTML(shuoming)
                    smdivs = reshuoming.xpath('//div')

                    for diai in range(len(smdivs)):
                        nrxx = etree.tostring(smdivs[diai], encoding="utf-8", pretty_print=True, method="html").decode("utf-8")
                        response = etree.HTML(nrxx)
                        chun = str(response.xpath('string(.)')).replace("\r\n", "").replace("\n", "")
                        ja = doc.add_paragraph()
                        r = ja.add_run("\n" + jm(chun, xxuid))
                        r.font.name = u'宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        r.font.size = Pt(11)
                        ja.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        tupians = response.xpath('//img[contains(@class, "ar_image_")]/@src')
                        # 这里区分是否存在图片
                        if (len(tupians) > 0) and (tupians != "[]"):
                            for tuim in range(len(tupians)):
                                response = requests.get(tupians[tuim])
                                binary_img = BytesIO(response.content)
                                imja = doc.add_picture(binary_img)
                                imja.alignment = WD_ALIGN_PARAGRAPH.CENTER


def wword(bian, biaotic, bianmac, biaojic, doc):
    '''
    global biaoti
    global biaozhunzhuangtai
    global mingzi
    global bianhao
    '''

    biaotia = ""
    cookies = {
        'rememberMe': 'YKYXKXwEaRou/r7CKSnlwGMFNNoq18/q0MzJo4lRKZjtjKVStCeWgOj3aYHb7exJ7TM7yzPZHIaBaUrNGDDdJnoOq0kxVkWjgbSVXOg0TFX7wGece61PmQaaB3zlq6RFISLrenLldZcjuMUkkzVHxxMa48BOrfYV9jTdfJD8CiRxqR0C1Pur1J7c6Ogt8WbBhFYhLHgVcVVDtxyyCmqBxvE70H8Fmfbq2k2p76Qwj7CqJuyIc+K3TqFjC5Pn50BZpwyS6KQvwH4aQisp5FTPU7G+7CoacsfJoznP8T3gmZBcWc14p8bqgXSnHwMQJNDe66T1YLTlo8juhWzWjUhsa6s7hiGQU85kxVnIko0i92rTZVF5OWRJE5ItW2OeqgQ6wnmH6AMIHGDfCmogiOA1nrMQsBhtBG0cMxHSnPBiwTSb1/j7ZvMqE/JaSVuzpqI+O6OdqLxm+N3gA7UagkidhcgDmjQjqCySCHbmSnAnaGDHssTXyLUjaDyWEbG5ilRyud/EvgjaLU1BxPrqE5hdx543jvsgslIHO98C2aMQhL5qb1YwCthy3qK4vbZ/R2QvBLzA00J/6QHABjJI02nC1Sir',
        'JSESSIONID': 'B08BADCC09CDE541FD14E8C7AB9063C2',
        'JSESSIONID': 'B08BADCC09CDE541FD14E8C7AB9063C2',
        'clientlanguage': 'und',
    }

    headers = {
        'Connection': 'keep-alive',
        'Accept': '*/*',
        'X-Requested-With': 'XMLHttpRequest',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36',
        'Content-Type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'Origin': 'https://www.gongbiaoku.com',
        'Sec-Fetch-Site': 'same-origin',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': 'https://www.gongbiaoku.com/read/a2z18467ck3?secId=44j80245592y',
        'Accept-Language': 'zh-CN,zh;q=0.9',
        # Requests sorts cookies= alphabetically
        # 'Cookie': 'rememberMe=YKYXKXwEaRou/r7CKSnlwGMFNNoq18/q0MzJo4lRKZjtjKVStCeWgOj3aYHb7exJ7TM7yzPZHIaBaUrNGDDdJnoOq0kxVkWjgbSVXOg0TFX7wGece61PmQaaB3zlq6RFISLrenLldZcjuMUkkzVHxxMa48BOrfYV9jTdfJD8CiRxqR0C1Pur1J7c6Ogt8WbBhFYhLHgVcVVDtxyyCmqBxvE70H8Fmfbq2k2p76Qwj7CqJuyIc+K3TqFjC5Pn50BZpwyS6KQvwH4aQisp5FTPU7G+7CoacsfJoznP8T3gmZBcWc14p8bqgXSnHwMQJNDe66T1YLTlo8juhWzWjUhsa6s7hiGQU85kxVnIko0i92rTZVF5OWRJE5ItW2OeqgQ6wnmH6AMIHGDfCmogiOA1nrMQsBhtBG0cMxHSnPBiwTSb1/j7ZvMqE/JaSVuzpqI+O6OdqLxm+N3gA7UagkidhcgDmjQjqCySCHbmSnAnaGDHssTXyLUjaDyWEbG5ilRyud/EvgjaLU1BxPrqE5hdx543jvsgslIHO98C2aMQhL5qb1YwCthy3qK4vbZ/R2QvBLzA00J/6QHABjJI02nC1Sir; JSESSIONID=B08BADCC09CDE541FD14E8C7AB9063C2; JSESSIONID=B08BADCC09CDE541FD14E8C7AB9063C2; clientlanguage=und',
    }

    params = {
        'secId': 'bae900204nwu',
    }

    # print("读取内容的链接：https://www.gongbiaoku.com/read/"+bian)
    response = requests.get('https://www.gongbiaoku.com/read/' + bian, params=params, cookies=cookies, headers=headers)
    html = etree.HTML(response.text)

    html2 = ''
    bm = 0
    for i in html.xpath('.//div[1][@class="catalog-list"]/ul[1]/li'):
        biaoti = str(i.xpath('.//h1[1]/a[1]/text()')[0]).replace("	", "")
        xxuids = str(i.xpath('.//h1[1]/a[1]/@data-content-uuid'))
        if len(xxuids) > 0:
            xxuid = xxuids[0].replace("	", "")
        if (len(i.xpath('.//h1[1]/a[1]/@data-content-uuid')) == 0):
            html2 = html2 + '<h1>' + i.xpath('.//h1[1]/a[1]/text()')[0].strip() + '</h1>' + '<br></br><p></p>'
            pass
        else:
            xxuid = i.xpath('.//h1[1]/a[1]/@data-content-uuid')[0]
            html2 = html2 + '<h1>' + get_val3(i.xpath('.//h1[1]/a[1]/@data-content-uuid')[0]).split('<h2>')[1] + ''
            html2 = html2.replace(i.xpath('.//h1[1]/a[1]/text()')[0].strip() + '<button',
                                  i.xpath('.//h1[1]/a[1]/text()')[0].strip() + '</h1><p></p><br></br><h2><button')
        try:
            ismaa = 1
            for j in i.xpath('.//ul[@class="ch"]/li/h2[1]/a[1]'):
                xxuid = j.xpath('@data-content-uuid')[0]
                ismaa = 2
                html2 = html2 + '<h2>' + get_val3(j.xpath('@data-content-uuid')[0]).split('<h2>')[1] + '<p></p><br></br>'
                tbianhao = str(biaoti.replace("　", "").replace("\r\n", ""))
                doc.add_page_break()

                if biaotia != tbianhao:
                    biaotia = tbianhao
                    ja = doc.add_heading()
                    r = ja.add_run(tbianhao)
                    r.font.name = u'MS Gothic'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
                    r.font.size = Pt(14)
                    r.bold = True
                    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    ja = doc.add_paragraph()
                    r = ja.add_run()
                    r.font.name = u'MS Gothic'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
                    r.font.size = Pt(14)
                    r.bold = True
                    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 以《h2》为分割点
                h2split = html2.split("<h2>")
                isbj = 1

                if ("引用标准名录" in biaoti) or ("本规范用词说明" in biaoti):
                    #  zhl("<h2>" + h2split[imad],biaoti)
                    if (len(h2split) == 1):
                        zhl(html2.replace("<h1>", "<h2>"), biaoti, xxuid, doc)
                    else:
                        for imad in range(len(h2split)):
                            if imad > 0:
                                if isbj == 1:
                                    isbj = isbj + 1
                                    zhl("<h2>" + h2split[imad], biaoti, xxuid, doc)
                                else:
                                    doc.add_page_break()
                                    zhl("<h2>" + h2split[imad], biaoti, xxuid, doc)
                else:
                    if (len(h2split) == 1):
                        xiedoc(html2.replace("<h1>", "<h2>"), biaoti, xxuid, doc)
                    else:
                        for imad in range(len(h2split)):
                            if imad > 0:
                                if isbj == 1:
                                    isbj = isbj + 1
                                    xiedoc("<h2>" + h2split[imad], biaoti, xxuid, doc)
                                else:
                                    # 问题可能在这个位置上
                                    doc.add_page_break()
                                    xiedoc("<h2>" + h2split[imad], biaoti, xxuid, doc)
                html2 = ""
                pass
        except:
            pass

        if ismaa == 1:
            # bm 是为了 方便区分 位置   bm-0是第一个   bm=1 是前言之前包括前言
            if bm == 0:
                htmla = etree.HTML(html2)
                title = jm(str(htmla.xpath('//h1')[0].text), xxuid).replace(" ", "").split("公告")
                ja = doc.add_paragraph()
                r = ja.add_run(title[0] + "公告")
                r.font.name = u'宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                r.font.size = Pt(11)
                r.bold = True
                ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ja = doc.add_paragraph()

                if title[1].replace("\n", "").replace(" ", "") != "":
                    r = ja.add_run(title[1])
                    r.font.name = u'宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    r.font.size = Pt(11)
                    r.bold = True
                    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 正则表达式提取
                divas = re.findall(r"<div.*?</div>", html2)
                for im in range(len(divas)):
                    if 'div class="align-right"' in divas[im]:
                        response = etree.HTML(text=divas[im])
                        nr = response.xpath('string(.)')
                        ja = doc.add_paragraph()
                        r = ja.add_run(jm(nr, xxuid))
                        r.font.name = u'宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        r.font.size = Pt(11)
                        r.bold = True
                        ja.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif 'class="bold"' in divas[im]:
                        response = etree.HTML(text=divas[im])
                        nr = response.xpath('string(.)')
                        ja = doc.add_paragraph()
                        r = ja.add_run(jm(nr, xxuid))
                        r.font.name = u'宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        r.font.size = Pt(11)
                        r.bold = True
                        ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        response = etree.HTML(text=divas[im])
                        nr = response.xpath('string(.)')
                        ja = doc.add_paragraph()
                        r = ja.add_run(jm(nr.replace("        ", "     "), xxuid))
                        r.font.name = u'宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        r.font.size = Pt(11)
                        ja.alignment = WD_ALIGN_PARAGRAPH.LEFT
                bm = bm + 1
            elif bm == 1:
                htmla = etree.HTML(html2)
                title = jm(str(htmla.xpath('//h1')[0].text), xxuid).replace(" ", "")

                if title != "局部修订说明":
                    bm = bm + 1

                ja = doc.add_paragraph()
                r = ja.add_run("")
                r.font.name = u'宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                r.font.size = Pt(11)
                ja.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # ja = doc.add_paragraph()
                ja = doc.add_heading()
                ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = ja.add_run(title)
                r.font.name = u'宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                r.font.size = Pt(11)
                r.bold = True
                r.font.color.rgb = RGBColor(0, 0, 0)
                # 正则表达式提取
                divas = re.findall(r"<div.*?</div>", html2)
                ma = 1

                for im in range(len(divas)):
                    response = etree.HTML(text=divas[im])
                    nr = jm(str(response.xpath('string(.)')), xxuid)
                    ja = doc.add_paragraph()
                    # r = ja.add_run("    " + str(nr).replace("        ", ""))

                    if len(str(nr.replace(" ", ""))) > 0:
                        if str(nr.replace(" ", "")[0]).isdigit():
                            r = ja.add_run("\n     " + str(nr).replace("        ", ""))
                    else:
                        r = ja.add_run("    " + str(nr).replace("        ", ""))

                    r.font.name = u'宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    r.font.size = Pt(11)
                    ja.alignment = WD_ALIGN_PARAGRAPH.LEFT

            else:
                tbianhao = str(biaoti.replace("　", "").replace("\r\n", ""))

                # 用于区分 是否 写标题
                if biaotia != tbianhao:
                    biaotia = tbianhao
                    doc.add_page_break()
                    ja = doc.add_heading()
                    r = ja.add_run(tbianhao)
                    r.font.name = u'MS Gothic'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
                    r.font.size = Pt(14)
                    r.bold = True
                    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r.font.color.rgb = RGBColor(0, 0, 0)
                    ja = doc.add_paragraph()
                    r = ja.add_run()
                    r.font.name = u'MS Gothic'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
                    r.font.size = Pt(14)
                    r.bold = True
                    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 以《h2》为分割点
                h2split = html2.split("<h2>")
                isbj = 1

                if ("引用标准名录" in biaoti) or ("本规范用词说明" in biaoti):
                    if (len(h2split) == 1):
                        zhl(html2.replace("<h1>", "<h2>"), biaoti, xxuid, doc)
                    else:
                        for imad in range(len(h2split)):
                            if imad > 0:
                                if isbj == 1:
                                    isbj = isbj + 1
                                    zhl("<h2>" + h2split[imad], biaoti, xxuid, doc)
                                else:
                                    doc.add_page_break()
                                    zhl("<h2>" + h2split[imad], biaoti, xxuid, doc)
                else:
                    if (len(h2split) == 1):
                        xiedoc(html2.replace("<h1>", "<h2>"), biaoti, xxuid, doc)
                    else:
                        for imad in range(len(h2split)):
                            if imad > 0:
                                if isbj == 1:
                                    isbj = isbj + 1
                                    xiedoc("<h2>" + h2split[imad], biaoti, xxuid, doc)
                                else:
                                    doc.add_page_break()
                                    xiedoc("<h2>" + h2split[imad], biaoti, xxuid, doc)

        html2 = ""
        # 生成word的文件 保存
        book_name = str(bianmac).replace(" ", "") + str(biaotic).replace(" ", "") + str(biaojic).replace(" ", "")
        doc.save(book_name.replace("/", "-") + ".docx")
        # doc.save(str(bianhao).replace(" ","")+str(mingzi).replace(" ","")+str(biaozhunzhuangtai).replace(" ","")+".docx")


# 获取详情页面的信息
def xingqing(url):
    global gfm
    '''
    global conn
    global cur
    global doc
    global biaozhunzhuangtai
    global mingzi
    global bianhao
    '''

    doc = Document()

    headers = {
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.0.0 Safari/537.36'}

    # print("读取内容的链接：https://www.gongbiaoku.com/book/" + url)
    request = urllib.request.Request(url="https://www.gongbiaoku.com/book/" + url, headers=headers)
    response = urllib.request.urlopen(request)
    content = response.read().decode('utf-8')

    # 提取出对应的数据
    c_html = etree.HTML(content)

    # 中文名字
    mingzi = c_html.xpath('//div[@class="intr fl"]/h2')[0].text
    gfm["mingzi"] = mingzi

    # 英文名字
    mingziy = c_html.xpath('//div[@class="intr fl"]/h4')[0].text

    # 标准号
    ulli = c_html.xpath('//div[@class="intr fl"]/ul/li')
    bianhao = str(etree.tostring(ulli[0], method='html'), "utf-8").split('</span>')[1].split('</li>')[0]
    gfm["bianhao"] = bianhao
    bzzt = etree.HTML("<li>" + str(etree.tostring(ulli[2], method='html'), "utf-8").split('</span>')[1].split('</li>')[0])
    biaozhunzhuangtai = bzzt.xpath('//li')[0].text
    gfm["biaozhunzhuangtai"] = biaozhunzhuangtai

    # 实施日期
    shishiriqi = str(etree.tostring(ulli[3], method='html'), "utf-8").split('</span>')[1].split('</li>')[0].split('-')
    ssriqi = shishiriqi[0] + "年" + str(int(shishiriqi[1])) + "月" + str(int(shishiriqi[2])) + "日"
    zhubianbumen = str(c_html.xpath('//div[@class="content"]/ul/li/div')[0].text).replace(" ", "").replace("\r\n", "")

    divp = c_html.xpath('//div[@class="content"]/ul/li/div/p')

    # 这个位置因为后来修改 好像用不到了
    global sh
    if len(divp) < 6:
        sh["zhubianbumen"] = str(zhubianbumen)
        sh["pizhunbumen"] = str(divp[0].text)
        sh["fabudanwei"] = str(divp[1].text)
        sh["zhubiandanwei"] = str(divp[2].text)
        sh["canbiandanwei"] = str(divp[3].text)
        sh["zhuyaoqicaoren"] = str(divp[4].text)
    else:
        # 这个位置因为后来修改 好像用不到了
        sh["zhubianbumen"] = str(zhubianbumen)
        sh["pizhunbumen"] = str(divp[0].text)
        sh["fabudanwei"] = str(divp[1].text)
        sh["zhubiandanwei"] = str(divp[2].text)
        sh["canbiandanwei"] = str(divp[3].text)
        sh["zhuyaoqicaoren"] = str(divp[4].text)
        sh["zhuyaoshencharen"] = str(divp[5].text)

    # 下面是生成 word 的头部的
    paragraph = doc.add_heading()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(mingzi)
    run.font.name = u'MS Gothic'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'MS Gothic')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run = paragraph.add_run(bianhao.replace(" ", ""))
    run.font.name = u'Calibri'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    ja = doc.add_paragraph()
    r = ja.add_run("中华人民共和国国家标准")
    r.bold = True
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ja = doc.add_paragraph()
    r = ja.add_run(mingzi)
    r.bold = True
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ja = doc.add_paragraph()
    r = ja.add_run(mingziy)
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ja = doc.add_paragraph()
    r = ja.add_run(bianhao)
    r.bold = True
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 主编部门
    ja = doc.add_paragraph()
    r = ja.add_run("主编部门：" + zhubianbumen)
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 批准部门
    ja = doc.add_paragraph()
    r = ja.add_run("批准部门：" + divp[0].text)
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 施行日期
    ja = doc.add_paragraph()
    r = ja.add_run("施行日期：" + ssriqi)
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 空行
    ja = doc.add_paragraph()
    r = ja.add_run(" ")
    r.font.name = u'宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    r.font.size = Pt(11)
    ja.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r.font.color.rgb = RGBColor(0, 0, 0)

    # 调用内容函数
    wword(url, bianhao, mingzi, biaozhunzhuangtai, doc)


# 测试可以用这里的，直接填写上编号 就可以生成了
xingqing("a2z18467ck3")


# 列表
def listb():
    global gfm
    # xpath
    for ima in range(2):
        request = urllib.request.Request(url="https://www.gongbiaoku.com/search?pageNo=" + str(
            ima + 1) + "&query=&status=&itemCatIds=1190&orderField=top&asc=0&style=")
        response = urllib.request.urlopen(request)
        content = response.read().decode('utf-8')

        htmllist = etree.HTML(content)
        hreflist = htmllist.xpath('//ul[@class="box-list"]/li/a/@href')

        for ilist in range(len(hreflist)):
            print("正在执行的，链接：https://www.gongbiaoku.com" + str(hreflist[ilist]))
            xingqing(str(hreflist[ilist]).replace("/book/", ""))
            print(gfm)  # 这个就是刚下载完的

# listb()
